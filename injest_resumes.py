import os
import re
import pandas as pd
import tiktoken
import requests
import json
from dotenv import load_dotenv
import pyodbc
import struct
from azure.identity import DefaultAzureCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import docx


# Load environment variables
load_dotenv()

# Azure Document Intelligence setup
endpoint = os.getenv("AZUREDOCINTELLIGENCE_ENDPOINT")
api_key = os.getenv("AZUREDOCINTELLIGENCE_API_KEY")
document_analysis_client = DocumentAnalysisClient(
    endpoint=endpoint,
    credential=AzureKeyCredential(api_key)
)

def get_pdf_files(folder_path):
    for path, subdirs, files in os.walk(folder_path):
        for name in files:
            if name.endswith(".pdf"):
                yield os.path.join(path, name)

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        poller = document_analysis_client.begin_analyze_document("prebuilt-layout", document=f)
    result = poller.result()
    text = ""
    for page in result.pages:
        for line in page.lines:
            text += line.content + " "
    return text

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    return text

def split_text_into_token_chunks(text, max_tokens=500):
    tokenizer = tiktoken.get_encoding("cl100k_base")
    tokens = tokenizer.encode(text)
    chunks = []
    for i in range(0, len(tokens), max_tokens):
        chunk_tokens = tokens[i:i + max_tokens]
        chunk_text = tokenizer.decode(chunk_tokens)
        chunks.append(chunk_text)
    return chunks

def extract_text_from_docx(docx_file):
    """
    Extracts and returns all text from a DOCX file.
    Accepts either a file path or a file-like object.
    """
    if isinstance(docx_file, str):
        doc = docx.Document(docx_file)
    else:
        doc = docx.Document(docx_file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def main():
    # Path to the directory containing PDF files
    folder_path = os.path.join(os.getcwd(), "IT")
    pdf_files = [f for f in get_pdf_files(folder_path)]
    num_files = len(pdf_files)
    print(f"Number of PDF files in the directory: {num_files}")

    data = []
    for file_id, pdf_file in enumerate(pdf_files):
        file_name = os.path.basename(pdf_file)
        print(f"Processing file {file_id + 1}/{num_files}: {file_name}")
        pdf_path = os.path.join(folder_path, file_name)
        text = extract_text_from_pdf(pdf_path)
        cleaned_text = clean_text(text)
        chunks = split_text_into_token_chunks(cleaned_text)
        print(f"Number of chunks for file {file_name}: {len(chunks)}")
        for chunk_id, chunk in enumerate(chunks):
            chunk_text = chunk.strip() if chunk.strip() else "NULL"
            unique_chunk_id = f"{file_id}_{chunk_id}"
            print(f"File: {file_name}, Chunk ID: {chunk_id}, Unique Chunk ID: {unique_chunk_id}, Chunk Length: {len(chunk_text)}, Chunk Text: {chunk_text[:50]}...")
            data.append({
                "file_name": file_name,
                "chunk_id": chunk_id,
                "chunk_text": chunk_text,
                "unique_chunk_id": unique_chunk_id
            })

    df = pd.DataFrame(data)
    print("Columns in df:", df.columns.tolist())
    if not df.empty and 'chunk_text' in df.columns:
        df['chunk_length'] = df['chunk_text'].apply(len)
        print(df[['file_name', 'chunk_id', 'chunk_length']].head(5))
    else:
        print("DataFrame is empty or missing 'chunk_text' column. Check data extraction logic.")

    # PART 2: Generating Embeddings for Text Chunks using Azure Open AI

    openai_embedding_model = os.getenv("AZOPENAI_EMBEDDING_MODEL_DEPLOYMENT_NAME")
    openai_url = os.getenv("AZOPENAI_ENDPOINT") + "openai/deployments/" + openai_embedding_model + "/embeddings?api-version=2023-05-15"
    openai_key = os.getenv("AZOPENAI_API_KEY")

    def get_embedding(text):
        response = requests.post(openai_url,
            headers={"api-key": openai_key, "Content-Type": "application/json"},
            json={"input": [text]}
        )
        if response.status_code == 200:
            response_json = response.json()
            embedding = json.loads(str(response_json['data'][0]['embedding']))
            return embedding
        else:
            return None

    all_filenames = []
    all_chunkids = []
    all_chunks = []
    all_embeddings = []

    for index, row in df.iterrows():
        filename = row['file_name']
        chunkid = row['unique_chunk_id']
        chunk = row['chunk_text']
        embedding = get_embedding(chunk)
        if embedding is not None:
            all_filenames.append(filename)
            all_chunkids.append(chunkid)
            all_chunks.append(chunk)
            all_embeddings.append(embedding)
        if (index + 1) % 50 == 0:
            print(f"Completed {index + 1} rows")

    result_df = pd.DataFrame({
        'filename': all_filenames,
        'chunkid': all_chunkids,
        'chunk': all_chunks,
        'embedding': all_embeddings
    })

    print(result_df.head(5))

    # PART 3: Using Azure SQL DB as a Vector Database to store and query embeddings

    def get_mssql_connection():
        entra_connection_string = os.getenv('ENTRA_CONNECTION_STRING')
        sql_connection_string = os.getenv('SQL_CONNECTION_STRING')
        if entra_connection_string:
            credential = DefaultAzureCredential(exclude_interactive_browser_credential=False)
            token = credential.get_token('https://database.windows.net/.default')
            token_bytes = token.token.encode('UTF-16LE')
            token_struct = struct.pack(f'<I{len(token_bytes)}s', len(token_bytes), token_bytes)
            SQL_COPT_SS_ACCESS_TOKEN = 1256
            conn = pyodbc.connect(entra_connection_string, attrs_before={SQL_COPT_SS_ACCESS_TOKEN: token_struct})
        elif sql_connection_string:
            conn = pyodbc.connect(sql_connection_string)
        else:
            raise ValueError("No valid connection string found in the environment variables.")
        return conn

    conn = get_mssql_connection()
    cursor = conn.cursor()
    cursor.fast_executemany = True

    for index, row in result_df.iterrows():
        chunkid = row['chunkid']
        filename = row['filename']
        chunk = row['chunk']
        embedding = row['embedding']
        query = """
        INSERT INTO dbo.resumedocs (chunkid, filename, chunk, embedding)
        VALUES (?, ?, ?,CAST(CAST(? as NVARCHAR(MAX)) AS VECTOR(1536)))"""
        cursor.execute(query, chunkid, filename, chunk, json.dumps(embedding))

    conn.commit()
    print("Data inserted successfully into the 'resumedocs' table.")
    conn.close()

if __name__ == "__main__":
    main()